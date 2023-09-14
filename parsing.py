import requests
from bs4 import BeautifulSoup
from time import sleep
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


tickets_doc = docx.Document()
answer_doc = docx.Document()


def test_el(url_path):

    response = requests.get(url_path)
    soup = BeautifulSoup(response.text, 'lxml')
    question_answer = []
    data = soup.find_all('div', class_='question row')

    for i, dat in enumerate(data):
        question_number = dat.find('div', class_='question__number').text
        question = dat.find('div', class_='question__text').text
        title = tickets_doc.add_paragraph().add_run('\n\n\n{0}: {1}\n'.format(question_number, question))
        title.italic = True
        title.bold = True
        answer_data = dat.find_all('div', class_='question__answers-list-item')
        for j, answ in enumerate(answer_data):
            answer = answ.find('span', class_='label').text
            tickets_doc.add_paragraph(f'{j+1} - {answer}')
            if answ.find('input', value='true'):
                if i+1 == 10:
                    question_answer.append(f'{i + 1} - {j + 1}')
                else:
                    question_answer.append(f'{i+1} - {j+1},   ')
    answer_doc.add_paragraph(question_answer)


def run_test_el(count_page, url_path):
    for i in range(count_page):
        sleep(1)
        print('Билет №{}'.format(i+1))
        answer_doc.add_paragraph('Билет №{}'.format(i+1))
        url = url_path + str(i+1)
        ticket = tickets_doc.add_paragraph('\n\n\n' + '*'*10 + 'Билет №{}'.format(i+1) + '*'*10)
        ticket.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        test_el(url)


run_test_el(30, 'https://prombez24.com/ticket/?testId=208&ticketNum=')
tickets_doc.save('Экзаменационные билеты.docx')
answer_doc.save('Ответы на экзаменационные вопросы.docx')
print('Файлы записаны!')
